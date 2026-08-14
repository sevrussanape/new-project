from __future__ import annotations
from pathlib import Path
import re, subprocess, tempfile

SUPPORTED = {'.pdf','.docx','.txt','.md','.mp3','.wav','.m4a','.mp4','.mkv','.webm','.mov'}

def extract_docx(path):
    from docx import Document
    doc = Document(path)
    blocks=[]
    for p in doc.paragraphs:
        if p.text.strip(): blocks.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells=[c.text.strip() for c in row.cells]
            if any(cells): blocks.append(' | '.join(cells))
    return '\n'.join(blocks)

def extract_text_file(path):
    return Path(path).read_text(encoding='utf-8', errors='ignore')

def extract_pdf(path, ocr=True):
    import fitz
    doc=fitz.open(path); pages=[]
    for i,page in enumerate(doc):
        text=page.get_text('text').strip()
        if text:
            pages.append(f'[Page {i+1}]\n{text}')
        elif ocr:
            try:
                import pytesseract
                from PIL import Image
                pix=page.get_pixmap(matrix=fitz.Matrix(1.7,1.7), alpha=False)
                img=Image.frombytes('RGB',[pix.width,pix.height],pix.samples)
                ocr_text=pytesseract.image_to_string(img).strip()
                if ocr_text: pages.append(f'[Page {i+1} OCR]\n{ocr_text}')
            except Exception:
                pass
    return '\n\n'.join(pages)

def extract_media(path, whisper_size='small'):
    from app.processors.audio import transcribe
    return transcribe(path, whisper_size)

def youtube_to_text(url, whisper_size='small'):
    """Uses transcript when available; otherwise downloads audio and transcribes locally."""
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        video_id=re.search(r'(?:v=|youtu\.be/|shorts/)([\w-]{11})',url).group(1)
        data=YouTubeTranscriptApi().fetch(video_id)
        text=' '.join(x.text for x in data)
        if text.strip(): return text, 'YouTube transcript'
    except Exception:
        pass
    try:
        import yt_dlp
        with tempfile.TemporaryDirectory() as td:
            out=str(Path(td)/'audio.%(ext)s')
            opts={'format':'bestaudio/best','outtmpl':out,'quiet':True,'noplaylist':True}
            with yt_dlp.YoutubeDL(opts) as ydl: info=ydl.extract_info(url,download=True); path=ydl.prepare_filename(info)
            text=extract_media(path,whisper_size)
            return text, 'YouTube audio + Whisper'
    except Exception as e:
        raise RuntimeError(f'YouTube import failed: {e}')

def image_text(path):
    try:
        import pytesseract
        from PIL import Image
        return pytesseract.image_to_string(Image.open(path)).strip()
    except Exception as e:
        raise RuntimeError('OCR requires pytesseract/Pillow and a local Tesseract installation.') from e

def chunk_text(text, chunk_size=1800, overlap=250):
    text=' '.join(text.split())
    chunks=[]; start=0
    while start < len(text):
        end=min(len(text),start+chunk_size); chunk=text[start:end]
        if chunk.strip(): chunks.append(chunk)
        if end==len(text): break
        start=end-overlap
    return chunks
